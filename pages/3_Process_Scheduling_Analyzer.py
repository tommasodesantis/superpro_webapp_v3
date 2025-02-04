import streamlit as st
from utils.check_auth import check_auth

# Check authentication
check_auth()
import pandas as pd
import io
import os
from datetime import datetime
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page config
st.set_page_config(page_title="Scheduling Analyzer", page_icon="📅")

# Title and description
st.title("Process Scheduling Analysis Tool")
with st.expander("Instructions", expanded=False):
    st.markdown("""
    Optimize your process scheduling with data-driven insights! This tool analyzes your SuperPro Designer scheduling data to:

    - Identify and resolve process bottlenecks
    - Optimize equipment utilization
    - Improve resource allocation
    - Minimize downtime and conflicts
    - Enhance batch campaign planning

    The analysis covers:
    - Bottleneck identification and mitigation
    - Equipment utilization patterns
    - Resource allocation optimization
    - Risk assessment and mitigation
    - Cleaning and maintenance scheduling
    - Scale-up considerations
    - Campaign planning recommendations

    Supported file formats:
    - Excel files (.xls/.xlsx) containing scheduling information

    Upload your scheduling data file below to begin the analysis.

    
    **How to export scheduling data to MS Project:**
    
    1) Open the Gantt Chart:

    - Navigate to the Gantt chart you wish to export. You can do this by selecting **Tasks** ▶ **Gantt Charts** from the main menu. Choose either the **Operations Gantt Chart** or **Equipment Gantt Chart** as needed.
    
    2) Export the Data:

    - Once the Gantt chart is open, go to **File** ▶ **Export to MS Project**. This option will create an XML file formatted specifically for MS Project that can be saved as .xsl.                     
     """)

def init_requesty():
    """Initialize OpenAI client with Requesty base URL"""
    return openai.OpenAI(
        base_url="https://router.requesty.ai/v1",
        api_key=st.secrets["REQUESTY_API_KEY"],
        default_headers={
            "HTTP-Referer": "https://github.com/yourusername/yourrepo",
            "X-Title": "Process Scheduling Analysis Tool"
        }
    )

def parse_scheduling_data(content):
    """Parse the scheduling data text content and return formatted text."""
    sections = []
    current_section = []
    lines = content.split('\n')
    
    # Process each line
    for line in lines:
        # Check if line is empty without stripping
        if not line.strip():
            if current_section:  # If we have a section built up
                sections.append('\n'.join(current_section))
                current_section = []
            continue
            
        # Check for section separator without modifying the line
        if '****' in line:  # Section separator
            if current_section:  # If we have a section built up
                sections.append('\n'.join(current_section))
                current_section = []
            current_section.append(line)
        else:
            # Add the line exactly as is, preserving all whitespace and tabs
            current_section.append(line)
    
    # Add any remaining section
    if current_section:
        sections.append('\n'.join(current_section))
    
    # Format output text
    formatted_text = []
    
    # Process each section
    for section in sections:
        lines = section.split('\n')
        
        # Skip empty sections
        if not any(line.strip() for line in lines):
            continue
            
        # Add section to formatted text
        formatted_text.append('\n'.join(lines))
    
    return '\n\n'.join(formatted_text)

def create_docx(content):
    """Create a DOCX document with proper markdown formatting."""
    doc = Document()
    
    # Define styles
    styles = doc.styles
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15
    
    # Add title
    title = doc.add_paragraph("Process Scheduling Analysis Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # Process content
    lines = content.split('\n')
    current_paragraph = None
    list_level = 0
    in_code_block = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()  # Preserve leading whitespace
        
        # Skip empty lines but add paragraph breaks
        if not line:
            current_paragraph = None
            i += 1
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                current_paragraph = doc.add_paragraph()
                current_paragraph.style = 'No Spacing'
                current_paragraph.paragraph_format.space_before = Pt(6)
                current_paragraph.paragraph_format.space_after = Pt(6)
            i += 1
            continue
            
        if in_code_block:
            if current_paragraph:
                current_paragraph.add_run(line + '\n')
            else:
                current_paragraph = doc.add_paragraph(line + '\n')
            current_paragraph.style = 'No Spacing'
            i += 1
            continue
            
        # Handle headings
        if line.startswith('#'):
            level = len(line.split()[0])  # Count the number of #
            text = line.lstrip('#').strip()
            p = doc.add_paragraph(text)
            p.style = f'Heading {min(level, 3)}'
            p.paragraph_format.space_before = Pt(12)
            current_paragraph = None
            
        # Handle bullet points and nested lists
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            
            # Extract the actual content
            content = line.lstrip('- ').lstrip('* ').strip()
            
            if current_paragraph is None or current_paragraph.style.name != 'List Bullet' or list_level != level:
                current_paragraph = doc.add_paragraph(content)
                current_paragraph.style = 'List Bullet'
                # Adjust indentation for nested lists
                current_paragraph.paragraph_format.left_indent = Pt(18 * (level + 1))
                list_level = level
            else:
                current_paragraph.add_run('\n' + content)
            
        # Handle numbered lists
        elif line.lstrip().startswith(tuple(f"{n}." for n in range(1, 10))):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            
            # Extract the actual content
            content = line.lstrip('123456789.').strip()
            
            if current_paragraph is None or current_paragraph.style.name != 'List Number' or list_level != level:
                current_paragraph = doc.add_paragraph(content)
                current_paragraph.style = 'List Number'
                # Adjust indentation for nested lists
                current_paragraph.paragraph_format.left_indent = Pt(18 * (level + 1))
                list_level = level
            else:
                current_paragraph.add_run('\n' + content)
            
        # Regular paragraphs
        else:
            if current_paragraph is None or current_paragraph.style.name in ['List Bullet', 'List Number']:
                current_paragraph = doc.add_paragraph(line)
                current_paragraph.style = 'Normal'
                list_level = 0
            else:
                current_paragraph.add_run('\n' + line)
        
        i += 1
    
    return doc

# Main interface
# File upload
uploaded_file = st.file_uploader("Upload scheduling data file", type=['xls', 'xlsx'])

if uploaded_file:
    if st.button("Analyze Schedule"):
        with st.spinner("Analyzing scheduling data..."):
            # Read file content
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            # Read all files as text since the .xls file is actually a text file
            content = uploaded_file.getvalue().decode('utf-8', errors='replace')
            
            # Parse the content
            formatted_text = parse_scheduling_data(content)
            
            # Initialize Requesty client
            client = init_requesty()
            
            # Generate analysis
            st.text("Generating analysis...")
            analysis_prompt = f"""You will be analyzing a detailed process scheduling dataset from SuperPro Designer. The data is presented in a tabular format with multiple sections.

IMPORTANT: Format your response using proper markdown syntax:
- Use # for main headers (e.g., # 1. Bottleneck Analysis)
- Use ## for subheaders (e.g., ## 1.1 Current Bottlenecks)
- Use - for bullet points
- Do NOT use bold text (**) for headers
- Do NOT use numbered sections without # or ##
- Each section must start with a proper header using # or ##
- Start each main section with a brief overview paragraph

The dataset includes:
- Process Parameters section showing operating time, campaigns, and batch information
- Equipment and Procedure Overview table listing all procedures with their operating modes and equipment
- Detailed Procedure tables showing operations with timing information (setup, process, and turnover times)

Please analyze this structured data and provide insights and recommendations in the following areas:

# 1. Bottleneck Analysis
## 1.1 Current Bottlenecks
- Identify critical process bottlenecks
- Quantify impact on throughput
- Map bottleneck dependencies

## 1.2 Equipment Utilization
- Analyze equipment usage patterns
- Identify capacity constraints
- Evaluate scheduling conflicts

## 1.3 Mitigation Strategies
- Propose specific improvements
- Estimate impact of changes
- Prioritize interventions

# 2. Scheduling Optimization
## 2.1 Current Efficiency
- Evaluate scheduling patterns
- Identify inefficiencies
- Analyze cycle times

## 2.2 Improvement Opportunities
- Recommend scheduling adjustments
- Propose conflict resolution
- Suggest timing optimizations

# 3. Resource Utilization
## 3.1 Equipment Analysis
- Detail usage patterns
- Identify peak demands
- Map resource conflicts

## 3.2 Labor Requirements
- Analyze shift patterns
- Evaluate workload distribution
- Identify staffing needs

# 4. Risk Assessment
## 4.1 Scheduling Risks
- Identify critical paths
- Analyze failure impacts
- Map dependencies

## 4.2 Mitigation Planning
- Propose contingencies
- Recommend backups
- Outline recovery procedures

# 5. Maintenance Strategy
## 5.1 Current Schedule
- Review maintenance windows
- Analyze cleaning requirements
- Evaluate downtimes

## 5.2 Optimization Plan
- Propose improved schedules
- Suggest preventive measures
- Recommend monitoring

# 6. Scale-up Analysis
## 6.1 Capacity Assessment
- Evaluate current utilization
- Identify constraints
- Analyze growth potential

## 6.2 Scale-up Strategy
- Propose expansion path
- Identify critical upgrades
- Recommend sequence

# 7. Campaign Planning
## 7.1 Current Structure
- Analyze batch sequences
- Evaluate campaign efficiency
- Identify transitions

## 7.2 Optimization Plan
- Recommend improvements
- Propose sequence changes
- Suggest transition optimization

# 8. Summary and Recommendations
## 8.1 Key Findings
- Summarize critical issues
- Present major opportunities
- Highlight priorities

## 8.2 Implementation Plan
- Propose action sequence
- Outline resource needs
- Define success metrics

Data:
{formatted_text}

For each area, carefully examine the tabular data and provide detailed insights and recommendations. Reference specific procedures, equipment, and timing data from the tables to support your analysis. Pay special attention to the relationships between procedures, their timing, and equipment utilization patterns shown in the data tables."""

            response = client.chat.completions.create(
                model="cline/o3-mini",
                messages=[{"role": "user", "content": analysis_prompt}]
            )
            analysis_content = response.choices[0].message.content
            
            # Create document
            st.text("Formatting document...")
            doc = create_docx(analysis_content)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"Process_Scheduling_Analysis_{timestamp}.docx"
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)
            
            # Provide download button
            st.download_button(
                label="Download Analysis Report",
                data=doc_bytes,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            # Display preview
            st.markdown("### Analysis Preview")
            st.markdown(analysis_content)
