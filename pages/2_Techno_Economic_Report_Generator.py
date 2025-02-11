import streamlit as st
from utils.check_auth import check_auth

# Check authentication
check_auth()
import pandas as pd
import io
import os
from datetime import datetime
import openai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Page config
st.set_page_config(page_title="Techno-economic Report Generator", page_icon="💸")

# Title and description
st.title("Techno-economic Report Generator")
with st.expander("Instructions", expanded=False):
    st.markdown("""
    Transform your SuperPro Designer data into comprehensive techno-economic analysis reports! This tool:

    - Analyzes process design parameters and economic metrics
    - Generates detailed capital investment breakdowns
    - Calculates operating costs and key performance indicators
    - Provides data-driven recommendations for optimization

    Required files:
    1. **Input Data Report (IDR)**: Contains process parameters, equipment specifications, and material balances
    2. **Economic Evaluation Report (EER)**: Includes capital costs, operating expenses, and economic metrics

    The generated report will include:
    - Process overview and technical analysis
    - Detailed economic evaluation
    - Cost driver analysis
    - Profitability assessment
    - Optimization recommendations

    Upload your Excel files (.xlsx/.xls) below to begin the analysis.
    """)

def init_openrouter():
    """Initialize OpenAI client with OpenRouter base URL"""
    return openai.OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=st.secrets["OPENROUTER_API_KEY"],
        default_headers={
            "HTTP-Referer": "https://github.com/yourusername/yourrepo",
            "X-Title": "TEA Analysis Tool"
        }
    )

def init_requesty():
    """Initialize OpenAI client with Requesty base URL"""
    return openai.OpenAI(
        base_url="https://router.requesty.ai/v1",
        api_key=st.secrets["REQUESTY_API_KEY"],
        default_headers={
            "HTTP-Referer": "https://github.com/yourusername/yourrepo",
            "X-Title": "TEA Analysis Tool"
        }
    )

def process_excel_file(file):
    """Process Excel file and return selected sheets as dataframes."""
    try:
        # Read all sheets
        xls = pd.ExcelFile(file)
        sheets = xls.sheet_names
        
        if len(sheets) > 1:
            selected_sheets = st.multiselect(
                f"Select sheets to process from {file.name}:",
                sheets,
                default=sheets
            )
        else:
            selected_sheets = sheets
            
        return {sheet: pd.read_excel(file, sheet_name=sheet) for sheet in selected_sheets}
    except Exception as e:
        st.error(f"Error processing {file.name}: {str(e)}")
        return {}

def dataframe_to_text(df):
    """Convert DataFrame to plain text, excluding empty cells."""
    # Drop completely empty rows and columns
    df = df.dropna(how='all').dropna(axis=1, how='all')
    
    # Convert to string and replace NaN with empty string
    text_table = df.fillna('').astype(str)
    
    # Generate text representation
    rows = []
    # Add header
    headers = '\t'.join(str(col) for col in text_table.columns)
    rows.append(headers)
    
    # Add data rows
    for _, row in text_table.iterrows():
        # Only include non-empty cells
        row_values = [str(val) for val in row if str(val).strip() != '']
        if row_values:  # Only add row if it has non-empty values
            rows.append('\t'.join(row_values))
    
    return '\n'.join(rows)

def create_docx(content):
    """Create a DOCX document with proper markdown formatting."""
    doc = Document()
    
    # Define styles
    styles = doc.styles
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(12)
    normal_style.paragraph_format.line_spacing = 1.15
    
    # Add title
    title = doc.add_paragraph("Techno-Economic Analysis Report")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    def process_text_formatting(text):
        """Process bold and italic markdown within text."""
        parts = []
        current_text = ""
        in_bold = False
        in_italic = False
        i = 0
        
        while i < len(text):
            if text[i:i+2] == '**' and not in_italic:
                if in_bold:
                    parts.append((current_text, True, False))
                    current_text = ""
                    in_bold = False
                    i += 2
                else:
                    if current_text:
                        parts.append((current_text, False, False))
                    current_text = ""
                    in_bold = True
                    i += 2
            elif text[i:i+1] == '*' and not in_bold:
                if in_italic:
                    parts.append((current_text, False, True))
                    current_text = ""
                    in_italic = False
                    i += 1
                else:
                    if current_text:
                        parts.append((current_text, False, False))
                    current_text = ""
                    in_italic = True
                    i += 1
            else:
                current_text += text[i]
                i += 1
        
        if current_text:
            parts.append((current_text, in_bold, in_italic))
        
        return parts
    
    # Process content
    lines = content.split('\n')
    current_paragraph = None
    list_level = 0
    in_code_block = False
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()  # Preserve leading whitespace
        
        # Skip empty lines but add paragraph breaks
        if not line:
            current_paragraph = None
            i += 1
            continue
            
        # Handle code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                current_paragraph = doc.add_paragraph()
                current_paragraph.style = 'No Spacing'
                current_paragraph.paragraph_format.space_before = Pt(6)
                current_paragraph.paragraph_format.space_after = Pt(6)
            i += 1
            continue
            
        if in_code_block:
            if current_paragraph:
                current_paragraph.add_run(line + '\n')
            else:
                current_paragraph = doc.add_paragraph(line + '\n')
            current_paragraph.style = 'No Spacing'
            i += 1
            continue
            
        # Handle headings (including numbered sections)
        if line.startswith('#'):
            level = len(line.split()[0])  # Count the number of #
            text = line.lstrip('#').strip()
            p = doc.add_paragraph(text)
            p.style = f'Heading {min(level, 3)}'
            p.paragraph_format.space_before = Pt(12)
            current_paragraph = None
        elif any(line.startswith(f"{n}." ) for n in range(1, 10)) and ":" in line:
            # Handle numbered sections with colons (e.g., "2.1 Methods:")
            p = doc.add_paragraph(line)
            p.style = 'Heading 2'
            p.paragraph_format.space_before = Pt(12)
            current_paragraph = None
            
        # Handle bullet points and nested lists
        elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            
            # Extract the actual content
            content = line.lstrip('- ').lstrip('* ').strip()
            
            if current_paragraph is None or current_paragraph.style.name != 'List Bullet' or list_level != level:
                current_paragraph = doc.add_paragraph()
                current_paragraph.style = 'List Bullet'
                # Adjust indentation for nested lists
                current_paragraph.paragraph_format.left_indent = Pt(18 * (level + 1))
                list_level = level
                
                # Process formatting in bullet point content
                for text, is_bold, is_italic in process_text_formatting(content):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
            else:
                current_paragraph.add_run('\n')
                for text, is_bold, is_italic in process_text_formatting(content):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
            
        # Handle numbered lists
        elif line.lstrip().startswith(tuple(f"{n}." for n in range(1, 10))):
            indent = len(line) - len(line.lstrip())
            level = indent // 2
            
            # Extract the actual content
            content = line.lstrip('123456789.').strip()
            
            if current_paragraph is None or current_paragraph.style.name != 'List Number' or list_level != level:
                current_paragraph = doc.add_paragraph()
                current_paragraph.style = 'List Number'
                # Adjust indentation for nested lists
                current_paragraph.paragraph_format.left_indent = Pt(18 * (level + 1))
                list_level = level
                
                # Process formatting in numbered list content
                for text, is_bold, is_italic in process_text_formatting(content):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
            else:
                current_paragraph.add_run('\n')
                for text, is_bold, is_italic in process_text_formatting(content):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
            
        # Regular paragraphs
        else:
            if current_paragraph is None or current_paragraph.style.name in ['List Bullet', 'List Number']:
                current_paragraph = doc.add_paragraph()
                current_paragraph.style = 'Normal'
                list_level = 0
                
                # Process formatting in paragraph content
                for text, is_bold, is_italic in process_text_formatting(line):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
            else:
                current_paragraph.add_run('\n')
                for text, is_bold, is_italic in process_text_formatting(line):
                    run = current_paragraph.add_run(text)
                    run.bold = is_bold
                    run.italic = is_italic
        
        i += 1
    
    return doc

# Main interface
# File upload section with consistent spacing
st.subheader("Upload Files")
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown("##### Input Data Report (IDR)")
    idr_file = st.file_uploader("Upload IDR file", type=['xlsx', 'xls'], key="idr_uploader")
    
with col2:
    st.markdown("##### Economic Evaluation Report (EER)")
    eer_file = st.file_uploader("Upload EER file", type=['xlsx', 'xls'], key="eer_uploader")

st.markdown("---")  # Add a separator

if idr_file and eer_file:
    if st.button("Generate Report"):
        with st.spinner("Processing files and generating report..."):
                # Process IDR file
                st.text("Processing IDR file...")
                idr_data = process_excel_file(idr_file)
                idr_text = "\n\n".join(dataframe_to_text(df) for df in idr_data.values())
                
                # Process EER file
                st.text("Processing EER file...")
                eer_data = process_excel_file(eer_file)
                eer_text = "\n\n".join(dataframe_to_text(df) for df in eer_data.values())
                
                # Initialize OpenRouter client for technical sections
                client = init_openrouter()
                
                # Generate technical sections
                st.text("Generating technical analysis...")
                technical_prompt = f"""You are a professional technical writer specializing in bioprocess engineering. Based on the provided process design data, analyze the process and generate the introduction and methods sections of a techno-economic report.

IMPORTANT: Format your response using proper markdown syntax:
- Use # for main headers (e.g., # 1. Introduction)
- Use ## for subheaders (e.g., ## 1.1 Process Overview)
- Use - for bullet points
- Do NOT use bold text (**) for headers
- Do NOT use numbered sections without # or ##
- Each section must start with a proper header using # or ##

Data:
{idr_text}

Please analyze the data and generate the following sections using markdown formatting:

# Title
Generate a concise and relevant title for this techno-economic analysis.

# 1. Introduction
## 1.1 Process Overview
- Use relevant data to provide an overview of the process
- Describe the product and its applications
- Outline the scale of production

## 1.2 Process Description
- Detail the process steps and sections
- Describe key unit operations
- Explain process flow and integration

# 2. Methods
## 2.1 Chemical Components
- List main chemical components used in the process
- Specify key raw materials
- Detail important intermediates

## 2.2 Process Parameters
- Main operational data
- Key process conditions
- Critical control parameters

## 2.3 Economic Framework
- Main economic data
- Labor requirements
- Resource allocations"""

                technical_response = client.chat.completions.create(
                    model="google/gemini-2.0-flash-001",
                    messages=[{"role": "user", "content": technical_prompt}]
                )
                technical_sections = technical_response.choices[0].message.content
                
                # Initialize Requesty client for economic sections
                requesty_client = init_requesty()
                
                # Generate economic sections
                st.text("Generating economic analysis...")
                economic_prompt = f"""You are a professional technical writer specializing in bioprocess engineering. Based on the provided economic data, analyze the economic aspects and generate the remaining sections of a techno-economic report.

IMPORTANT: Format your response using proper markdown syntax:
- Use # for main headers (e.g., # 3. Results)
- Use ## for subheaders (e.g., ## 3.1 Capital Investment)
- Write in clear, flowing paragraphs
- Use data points from the provided information to support your analysis
- Each section must start with a proper header using # or ##
- Maintain a professional, analytical tone throughout
- Focus solely on analyzing the provided data

Previous Sections Context:
{technical_sections}

Data:
{eer_text}

Generate the following sections using markdown formatting:

# 3. Results

## 3.1 Capital Investment Analysis
## 3.1.1 Equipment Investment
Provide a comprehensive analysis of the equipment investment requirements. Detail the major equipment costs and their relative contributions to the total investment. Explain the significance of key equipment items and their associated costs. Include specific cost figures from the data to support your analysis.

## 3.1.2 Facility Investment
Present a detailed breakdown of facility investment components. Analyze the distribution of costs across building, construction, and infrastructure elements. Explain how these investments support the process operations. Include specific numbers from the data to illustrate the scale and importance of each investment category.

## 3.2 Operating Costs
## 3.2.1 Direct Costs
Analyze the direct operating costs, explaining how raw materials, labor, and utilities contribute to the overall operational expenses. Provide specific cost figures and consumption rates from the data. Explain the relationships between different cost components and their impact on the process economics.

## 3.2.2 Indirect Costs
Examine the indirect costs associated with facility operation and maintenance. Analyze how these costs relate to the direct operational expenses. Include specific figures from the data to demonstrate the relative importance of different indirect cost categories.

## 3.3 Key Performance Indicators
## 3.3.1 Production Metrics
Present a detailed analysis of the production-related performance indicators. Explain how production capacity relates to resource utilization and operational efficiency. Use specific figures from the data to illustrate the process performance.

## 3.3.2 Financial Metrics
Analyze the key financial indicators that demonstrate the project's economic performance. Explain the relationships between different financial metrics and what they reveal about the process economics. Include specific numbers from the data to support your analysis.

# 4. Discussion
## 4.1 Cost Analysis
Synthesize the findings from the capital and operating cost analyses. Identify and explain the major cost drivers based on the data. Evaluate how different cost components interact and influence overall economic performance.

## 4.2 Process Economics
Analyze the relationship between capital investments and operating costs. Examine how process parameters affect economic performance. Identify potential areas for cost optimization based on the analyzed data.

# 5. Conclusions and Recommendations
## 5.1 Key Findings
Summarize the most significant economic findings from the analysis. Present clear conclusions about the process economics based on the analyzed data. Focus on the most important metrics and their implications.

## 5.2 Recommendations
Provide specific, data-driven recommendations for improving economic performance. Prioritize suggestions based on their potential impact and feasibility. Focus on practical improvements that are supported by the economic analysis."""

                economic_response = requesty_client.chat.completions.create(
                    model="cline/o3-mini",
                    messages=[{"role": "user", "content": economic_prompt}]
                )
                
                # Combine sections
                full_report = technical_sections + "\n\n" + economic_response.choices[0].message.content
                
                # Create document
                st.text("Formatting document...")
                doc = create_docx(full_report)
                
                # Save document
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                output_filename = f"Techno_Economic_Analysis_Report_{timestamp}.docx"
                doc_bytes = io.BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                # Provide download button
                st.download_button(
                    label="Download Report",
                    data=doc_bytes,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # Display preview
                st.markdown("### Report Preview")
                st.markdown(full_report)
